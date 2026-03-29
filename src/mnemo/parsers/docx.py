"""DOCX parser — extracts text and metadata via python-docx."""

from __future__ import annotations

import hashlib
from pathlib import Path

from mnemo.parser import NoteDocument

try:
    from docx import Document as _Document  # type: ignore[import-untyped]

    _HAS_DOCX = True
except ImportError:
    _HAS_DOCX = False


def parse_docx(file_path: Path, vault_root: Path | None = None) -> NoteDocument:
    """Parse a DOCX file into a NoteDocument.

    Requires the ``python-docx`` library (``pip install python-docx>=1.0``).
    """
    if not _HAS_DOCX:
        raise ImportError(
            "python-docx is required for DOCX parsing. Install with: pip install 'mnemo-secondbrain[documents]'"
        )

    doc = _Document(str(file_path))

    # --- extract text and headings -------------------------------------------
    paragraphs: list[str] = []
    headings: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        paragraphs.append(text)
        # python-docx marks heading styles as "Heading 1", "Heading 2", etc.
        style_name = (para.style.name or "") if para.style else ""
        if style_name.startswith("Heading"):
            headings.append(text)

    body = "\n\n".join(paragraphs)

    # --- extract table text (appended after paragraphs) ----------------------
    table_lines: list[str] = []
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                table_lines.append(" | ".join(cells))
    if table_lines:
        body = body + "\n\n" + "\n".join(table_lines)

    # --- metadata / synthetic frontmatter ------------------------------------
    core = doc.core_properties
    title = (core.title or "").strip() or file_path.stem
    author = (core.author or "").strip()
    created = core.created.isoformat() if core.created else None

    frontmatter: dict = {
        "title": title,
        "type": "source",
        "format": "docx",
        "tags": ["docx", "external"],
    }
    if author:
        frontmatter["author"] = author
    if created:
        frontmatter["created"] = created

    # --- key ------------------------------------------------------------------
    if vault_root is not None:
        try:
            rel = file_path.resolve().relative_to(vault_root.resolve())
            key = rel.with_suffix("").as_posix()
        except ValueError:
            key = file_path.stem
    else:
        key = file_path.stem

    content = body
    checksum = hashlib.sha256(content.encode("utf-8")).hexdigest()[:16]

    return NoteDocument(
        path=file_path,
        name=file_path.stem,
        key=key,
        content=content,
        body=body,
        frontmatter=frontmatter,
        wiki_links=[],
        tags=frontmatter["tags"],
        headings=headings[:30],
        checksum=checksum,
    )
